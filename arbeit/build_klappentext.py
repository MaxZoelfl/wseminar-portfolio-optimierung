#!/usr/bin/env python3
"""
Baut den Klappentext als .docx nach den W-Seminar-Formalien aus arbeit/Klappentext.md.

FÜR EINSTEIGER — WAS MACHT DIESES SKRIPT?
Das kleinste der drei Build-Skripte: Es wandelt den einseitigen Klappentext
(Kurzvorstellung der Arbeit) von Markdown in Word um. Weil das Dokument so
einfach ist, kommt es ganz OHNE pandoc aus — die Markdown-Datei wird Zeile
für Zeile selbst gelesen und direkt mit python-docx zum Word-Dokument
zusammengesetzt. (Erklärungen zu python-docx: siehe build_docx.py.)

Formalia (W-Seminararbeit_Formales.pdf):
  - Times New Roman 12, Zeilenabstand 1,5
  - Ränder: links 3,5 cm, rechts/oben/unten 2,5 cm
  - Blocksatz (optional laut Formalia, hier verwendet)

Markdown-Konvention (die Mini-Übersetzungstabelle dieses Skripts):
  "# ..."  -> Titelzeile (fett, zentriert)
  "## ..." -> Untertitel/Arbeitstitel (kursiv, zentriert)
  sonstige nichtleere Zeile -> Textabsatz (Blocksatz)

Aufruf:  venv/bin/python arbeit/build_klappentext.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Dateipfade: Quelle (MD) und Ziel (OUT).
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD = os.path.join(ROOT, "arbeit", "Klappentext.md")
OUT = os.path.join(ROOT, "arbeit", "Klappentext.docx")


def set_font(run_or_style, name="Times New Roman"):
    """Setzt die Schriftart robust — zusätzlich zur bequemen Eigenschaft auch
    direkt in Words XML für alle Zeichensatz-Varianten (vgl. build_docx.py)."""
    run_or_style.font.name = name
    rpr = run_or_style.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), name)


def main():
    # Leeres Word-Dokument anlegen (Standardvorlage von python-docx):
    doc = Document()

    # Grundstil "Normal" auf die Formalia bringen: TNR 12, 1,5-zeilig.
    normal = doc.styles["Normal"]
    normal.font.size = Pt(12)
    set_font(normal)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Seitenränder nach Formalia:
    for sec in doc.sections:
        sec.left_margin = Cm(3.5)
        sec.right_margin = Cm(2.5)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)

    # Markdown Zeile für Zeile einlesen und je nach Präfix formatieren:
    lines = open(MD, encoding="utf-8").read().splitlines()
    for line in lines:
        s = line.strip()                 # Leerraum an den Rändern entfernen
        if not s:
            continue                     # Leerzeilen überspringen
        if s.startswith("## "):
            # Untertitel: kursiv, zentriert, mit Abstand danach.
            # (s[3:] schneidet die ersten 3 Zeichen "## " ab.)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(18)
            r = p.add_run(s[3:]); set_font(r); r.font.size = Pt(12); r.italic = True
        elif s.startswith("# "):
            # Titelzeile: fett, etwas größer, zentriert.
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(s[2:]); set_font(r); r.font.size = Pt(14); r.bold = True
        else:
            # Normaler Textabsatz im Blocksatz (beide Ränder bündig).
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(s); set_font(r); r.font.size = Pt(12)

    doc.save(OUT)
    # Kontrollausgabe: Wortzahl des Fließtexts (Überschriften ausgenommen).
    words = sum(len(l.split()) for l in lines if l.strip() and not l.strip().startswith("#"))
    print(f"✓ {OUT}")
    print(f"  Textabsätze gebaut | ~{words} Wörter Fließtext")


if __name__ == "__main__":
    main()
