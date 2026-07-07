#!/usr/bin/env python3
"""
Reproduzierbarer Bau der W-Seminararbeit als .docx aus arbeit/arbeit.md.

FÜR EINSTEIGER — WAS MACHT DIESES SKRIPT?
Das Manuskript der Seminararbeit wird als einfache Textdatei (arbeit.md im
"Markdown"-Format: # = Überschrift, ** = fett, $…$ = Formel) geschrieben —
das lässt sich gut versionieren und bearbeiten. Abgegeben werden muss aber
ein Word-Dokument mit strengen Formalia (Schriftart, Ränder, Titelblatt …).
Dieses Skript ist die automatische "Druckerei": Es verwandelt die Textdatei
per Knopfdruck in ein fertig formatiertes Word-Dokument. Vorteil: Nach jeder
Textänderung entsteht das .docx identisch neu — nichts muss (fehleranfällig)
von Hand in Word nachformatiert werden.

Verwendete Werkzeuge:
  - pandoc      : ein verbreitetes Kommandozeilen-Programm, das zwischen
                  Dokumentformaten übersetzt (hier: Markdown → Word).
                  Es wird per subprocess.run(...) aufgerufen, also so, als
                  würde man den Befehl selbst ins Terminal tippen.
  - python-docx : eine Python-Bibliothek, die Word-Dateien direkt bearbeiten
                  kann — für alles, was pandoc nicht kann (Titelblatt,
                  Seitenzahlen ab Seite 3, Abschnittswechsel).
  - Die mit "w:" beginnenden Namen (w:sectPr, w:fldChar, …) sind Elemente
    von Word's internem XML-Speicherformat (OOXML). python-docx bietet nicht
    für alles bequeme Funktionen, daher wird an einigen Stellen direkt auf
    dieser XML-Ebene gearbeitet ("OxmlElement" = neues XML-Element erzeugen,
    "qn" = übersetzt Kurznamen wie "w:type" in die von Word erwartete
    ausgeschriebene Form).

Pipeline (die vier Schritte unten):
  1. Pandoc-Referenz mit Formalia-Stilen erzeugen (TNR 12, 1,5-zeilig, Ränder,
     Fußnoten 10).
  2. Markdown aufbereiten: HTML-Kommentar entfernen, [[FN: ...]] -> ^[...]
     (Pandoc-Inline-Fußnoten).
  3. Pandoc: Markdown -> docx mit Inhaltsverzeichnis (--toc) und echten Fußnoten.
  4. Nachbearbeitung (python-docx): Titelblatt, Abschnittswechsel,
     Seitennummerierung (Titel + Inhaltsverz. ohne Nummer, Text ab S. 3),
     Schlusserklärungs-Platzhalter, Anhang.

Aufruf:  venv/bin/python arbeit/build_docx.py
Voraussetzung: pandoc auf dem PATH (z. B. /opt/homebrew/bin), python-docx.
"""
import re, copy, subprocess, sys, os

# ---- Dateipfade zentral definiert -----------------------------------------
# ROOT = Projektordner (zwei Ebenen über dieser Datei); Zwischendateien
# landen in /tmp (dem Wegwerf-Ordner des Systems), nur OUT ist das Endprodukt.
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD   = os.path.join(ROOT, "arbeit", "arbeit.md")          # Quelle: das Manuskript
REF  = "/tmp/ref.docx"                                    # Formatvorlage (Schritt 1)
BODY = "/tmp/arbeit_body.docx"                            # Pandoc-Rohfassung (Schritt 3)
PMD  = "/tmp/arbeit_pandoc.md"                            # aufbereitetes Markdown (Schritt 2)
OUT  = os.path.join(ROOT, "arbeit", "W-Seminararbeit.docx")  # das fertige Dokument

from docx import Document
from docx.shared import Pt, Cm                            # Pt = Schriftgröße in Punkt, Cm = Zentimeter
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def step1_reference():
    """Schritt 1: Formatvorlage ("reference.docx") mit den Formalia bauen.

    Pandoc übernimmt beim Konvertieren alle Absatz-Stile aus einer solchen
    Referenzdatei. Wir lassen uns zuerst Pandocs Standardvorlage geben und
    biegen dann deren Stile auf die W-Seminar-Formalia um: Times New Roman 12,
    Zeilenabstand 1,5, gestaffelte Überschriftsgrößen, Fußnoten in 10 pt,
    Ränder links 3,5 cm / sonst 2,5 cm.
    """
    # Pandocs eingebaute Standard-Referenzdatei nach /tmp schreiben lassen:
    subprocess.run(["pandoc", "-o", "/tmp/ref_default.docx",
                    "--print-default-data-file", "reference.docx"], check=True)
    doc = Document("/tmp/ref_default.docx")

    def set_font(style, name="Times New Roman", size=None, bold=None):
        """Setzt Schriftart/-größe/-fett für einen Word-Stil — und zwar
        robust: zusätzlich zur bequemen python-docx-Eigenschaft werden die
        Schriftnamen auch direkt im XML für alle Zeichensatz-Varianten
        (ascii/hAnsi/cs) hinterlegt, damit Word sie wirklich überall anwendet."""
        style.font.name = name
        rpr = style.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            rf.set(qn(a), name)
        if size is not None: style.font.size = Pt(size)
        if bold is not None: style.font.bold = bold

    S = doc.styles
    # "Normal" = der Stil des Fließtexts: TNR 12, 1,5-zeilig, keine Extra-Abstände.
    n = S["Normal"]; set_font(n, size=12)
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    n.paragraph_format.space_after = Pt(0); n.paragraph_format.space_before = Pt(0)
    names = [s.name for s in S]
    # Überschriften-Hierarchie: H1 = 15 pt, H2 = 13 pt, H3 = 12 pt, alle fett.
    for h, sz in (("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)):
        if h in names:
            st = S[h]; set_font(st, size=sz, bold=True)
            st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(6)
            st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    # Fußnotentext: 10 pt, einzeilig. (Der Stil heißt je nach Word-Sprache
    # anders — beide Namen probieren; try/except fängt den fehlenden ab.)
    for fn in ("Footnote Text", "Fußnotentext"):
        try:
            st = S[fn]; set_font(st, size=10)
            st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except KeyError:
            pass
    # Titel-/Untertitel-Stile (falls vorhanden) ebenfalls auf TNR setzen:
    for t, sz in (("Title", 22), ("Subtitle", 14)):
        try: set_font(S[t], size=sz)
        except KeyError: pass
    # Seitenränder nach Formalia:
    for sec in doc.sections:
        sec.left_margin = Cm(3.5); sec.right_margin = Cm(2.5)
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    doc.save(REF)


def step2_preprocess():
    """Schritt 2: Das Markdown-Manuskript für pandoc aufbereiten.

    Zwei kleine Umbauten per regulärem Ausdruck (re.sub = "suchen & ersetzen
    nach Muster"):
      a) Der HTML-Kommentar am Dateianfang (<!-- … -->, interne Notizen)
         soll nicht im Word-Dokument landen → entfernen.
      b) Historische Fußnoten-Marker im Eigenformat [[FN: Text]] werden in
         pandocs Fußnoten-Syntax ^[Text] übersetzt, aus der pandoc ECHTE
         Word-Fußnoten erzeugt. (Aktuell nutzt das Manuskript APA-Kurzbelege
         statt Fußnoten; die Ersetzung bleibt für Altbestände drin.)
    Rückgabe: Anzahl der gefundenen Fußnoten-Marker (für die Kontrollausgabe).
    """
    src = open(MD, encoding="utf-8").read()
    src = re.sub(r"^<!--.*?-->\s*", "", src, count=1, flags=re.DOTALL)
    n = len(re.findall(r"\[\[FN:", src))
    src = re.sub(r"\[\[FN:\s*(.*?)\]\]", r"^[\1]", src)
    open(PMD, "w", encoding="utf-8").write(src)
    return n


def step3_pandoc():
    """Schritt 3: pandoc konvertiert das aufbereitete Markdown nach Word.

    Bedeutung der Optionen:
      --reference-doc  : unsere Formatvorlage aus Schritt 1 verwenden
      --toc            : automatisches Inhaltsverzeichnis einfügen
                         (bis Gliederungstiefe 3)
      --resource-path  : wo pandoc eingebundene Bilder (die Abbildungen aus
                         output1.6/) suchen soll
      -f markdown+...  : Eingabeformat; "+tex_math_dollars" aktiviert die
                         $…$-Schreibweise für mathematische Formeln
    """
    subprocess.run(["pandoc", PMD, "-o", BODY, "--reference-doc=" + REF,
                    "--toc", "--toc-depth=3", "--resource-path=" + ROOT,
                    "-f", "markdown+tex_math_dollars"],
                   check=True)


def step4_postprocess():
    """Schritt 4: Feinschliff am Word-Dokument (alles, was pandoc nicht kann).

    Im Einzelnen: Titelblatt VOR das Inhaltsverzeichnis setzen, einen
    Abschnittswechsel vor Kapitel 1 einbauen (damit die Seitennummerierung
    erst dort beginnt), die Seitenzahl-Fußzeile anlegen sowie Platzhalter-
    Seiten für Schlusserklärung und Anhang anhängen.
    """
    doc = Document(BODY); body = doc.element.body
    paras = doc.paragraphs
    # Zwei Ankerpunkte im Dokument suchen: den ersten Absatz (dort beginnt
    # das von pandoc erzeugte Inhaltsverzeichnis — vor ihn kommt das
    # Titelblatt) und die Überschrift "1 Einleitung" (vor ihr beginnt der
    # nummerierte Textteil).
    toc_anchor = paras[0]
    einleitung = next(p for p in paras if p.style.name.startswith("Heading")
                      and p.text.strip().startswith("1 Einleitung"))

    def ins(anchor, text, *, size=12, bold=False, align="center", sa=0, italic=False):
        """Hilfsfunktion: fügt VOR dem Anker-Absatz eine Textzeile ein
        (Ausrichtung, Größe, fett/kursiv und Abstand danach einstellbar).
        Damit wird das Titelblatt Zeile für Zeile aufgebaut."""
        p = anchor.insert_paragraph_before(); p.style = doc.styles["Normal"]
        p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                       "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
        p.paragraph_format.space_after = Pt(sa); p.paragraph_format.space_before = Pt(0)
        if text:
            r = p.add_run(text); r.font.name = "Times New Roman"
            r.font.size = Pt(size); r.bold = bold; r.italic = italic
        return p

    # ---- Titelblatt (die [Klammer-Texte] sind bewusst Platzhalter, die der
    # Verfasser vor der Abgabe von Hand ausfüllt; Leerzeilen schieben die
    # Blöcke vertikal auseinander) ------------------------------------------
    ins(toc_anchor, "[Name der Schule]", sa=2)
    ins(toc_anchor, "W-Seminar im Fach [Fach] — [Titel des W-Seminars]")
    for _ in range(4): ins(toc_anchor, "")
    ins(toc_anchor, "Markowitz versus Random Forest", size=22, bold=True, sa=4)
    ins(toc_anchor, "Ein statistisch abgesicherter Vergleich von Portfolio­optimierungs­strategien",
        size=14, italic=True)
    for _ in range(5): ins(toc_anchor, "")
    ins(toc_anchor, "Seminararbeit", bold=True, sa=2)
    ins(toc_anchor, "von [Vorname Nachname]")
    for _ in range(4): ins(toc_anchor, "")
    ins(toc_anchor, "Betreuende Lehrkraft: [Name der Lehrkraft]", sa=2)
    ins(toc_anchor, "Abgabetermin: [TT.MM.JJJJ]", sa=2)
    ins(toc_anchor, "[Ort], [TT.MM.JJJJ]")
    # Seitenumbruch nach dem Titelblatt, dann Überschrift fürs Inhaltsverzeichnis:
    toc_anchor.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
    h = toc_anchor.insert_paragraph_before("Inhaltsverzeichnis"); h.style = doc.styles["Heading 1"]

    # ---- Abschnittswechsel vor Kapitel 1 -----------------------------------
    # Word gliedert Dokumente in "Abschnitte" (sections); nur an ihren Grenzen
    # kann die Fußzeile/Seitennummerierung wechseln. Hier wird die Abschnitts-
    # definition vom Dokumentende kopiert und vor "1 Einleitung" eingesetzt —
    # damit sind Titelblatt+Inhaltsverzeichnis Abschnitt 1 (ohne Seitenzahl)
    # und der Textteil Abschnitt 2 (mit Seitenzahl).
    final = body.find(qn("w:sectPr"))
    sect1 = copy.deepcopy(final)
    for t in sect1.findall(qn("w:type")): sect1.remove(t)
    typ = OxmlElement("w:type"); typ.set(qn("w:val"), "nextPage"); sect1.insert(0, typ)
    # Den letzten echten Absatz VOR der Einleitung finden und den
    # Abschnittswechsel in dessen Absatz-Eigenschaften einhängen:
    prev = einleitung._p.getprevious()
    while prev is not None and prev.tag != qn("w:p"): prev = prev.getprevious()
    prev.get_or_add_pPr().append(sect1)

    # ---- Seitenzahlen in der Fußzeile von Abschnitt 2 ----------------------
    # Die Seitenzahl ist in Word ein "Feld" (automatisch aktualisierter
    # Platzhalter). Ein Feld besteht im XML aus drei Teilen: Start-Marke,
    # Feldanweisung (" PAGE ") und End-Marke — genau die baut die Schleife.
    secs = doc.sections
    s2 = secs[1]; s2.footer.is_linked_to_previous = False   # eigene Fußzeile ab Textteil
    fp = s2.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    for kind in ("begin", "text", "end"):
        if kind == "text":
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
            it.text = " PAGE "; run._r.append(it)
        else:
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), kind); run._r.append(fc)
    run.font.name = "Times New Roman"; run.font.size = Pt(11)
    secs[0].footer.is_linked_to_previous = False   # Abschnitt 1 behält seine LEERE Fußzeile

    # ---- Schlusserklärung und Anhang (je auf neuer Seite) -------------------
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("Schlusserklärung").style = doc.styles["Heading 1"]
    doc.add_paragraph("[Hier fügt der Verfasser die unterschriebene Schlusserklärung gemäß den "
                      "Formalien ein (inkl. Angabe der verwendeten KI-Tools und ihres Verwendungszwecks).]")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("Anhang").style = doc.styles["Heading 1"]
    doc.add_paragraph("Der vollständige Quellcode der Implementierung, sämtliche erzeugten Abbildungen "
                      "(kumulierte Renditen, Gewichtsverläufe, Effizienzlinie, rollierender Sharpe, "
                      "Feature-Importance/SHAP, Stress-Test, Turnover-Analyse) sowie die Protokolldateien "
                      "(CSV/JSON) befinden sich im begleitenden Repository. Die genauen Wortlaute der "
                      "verwendeten KI-Prompts sind hier aufzuführen.")
    doc.save(OUT)
    return len(doc.sections), len(doc.tables)


# Hauptprogramm: die vier Schritte nacheinander ausführen und eine kurze
# Kontrollzeile drucken (✓ Pfad + Statistik).
if __name__ == "__main__":
    step1_reference()
    nfn = step2_preprocess()
    step3_pandoc()
    nsec, ntbl = step4_postprocess()
    print(f"✓ {OUT}")
    print(f"  Fußnoten: {nfn} | Abschnitte: {nsec} | Tabellen: {ntbl}")
